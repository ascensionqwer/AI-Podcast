"""
Podcast Generator Module.
Handles the complete pipeline from content to podcast audio.
"""

import logging
import re
import time
import numpy as np
from dataclasses import dataclass
from pathlib import Path
from typing import Optional, List, Tuple
from concurrent.futures import ThreadPoolExecutor, as_completed

from openai import OpenAI

from .config import Config, get_config
from .tts_server import EmbeddedTTSServer, is_port_in_use

logger = logging.getLogger(__name__)


@dataclass
class PodcastSegment:
    """A single segment of the podcast."""
    speaker: str
    text: str
    voice: str


@dataclass
class PodcastScript:
    """Generated podcast script."""
    segments: List[PodcastSegment]
    raw_text: str
    source_file: Optional[str] = None


@dataclass
class CoverageResult:
    """Result of script coverage generation."""
    segments: List[PodcastSegment]
    structured_report: str
    raw_text: str
    source_file: Optional[str] = None


class LLMClient:
    """Client for interacting with LM Studio (OpenAI-compatible API)."""
    
    def __init__(self, config: Config):
        self.config = config
        
        # Create custom HTTP client with very long timeout for large context processing
        # Large contexts (100k+ tokens) can take 15+ minutes to process
        import httpx
        http_client = httpx.Client(
            timeout=httpx.Timeout(
                connect=30.0,      # Connection timeout
                read=1800.0,       # Read timeout: 30 minutes
                write=1800.0,      # Write timeout: 30 minutes
                pool=30.0          # Pool timeout
            )
        )
        
        self.client = OpenAI(
            base_url=config.llm.base_url,
            api_key=config.llm.api_key,
            http_client=http_client  # Use custom client with extended timeout
        )
        self.model = config.llm.model
    
    def check_connection(self) -> bool:
        """Check if LM Studio is running and accessible."""
        try:
            # Try a minimal request
            self.client.models.list()
            return True
        except Exception as e:
            logger.error(f"Cannot connect to LM Studio at {self.config.llm.base_url}: {e}")
            return False
    
    def _calculate_word_count(self, content: str, podcast_mode: str) -> int:
        """
        Calculate word count based on podcast mode.
        
        Modes:
        - "summary": Quick overview (~2500 words = ~17 min) - key points with context
        - "analysis": Detailed discussion (~4000 words) - deep insights and analysis
        - "full": Complete coverage - covers EVERYTHING, dynamic based on content
        
        Args:
            content: The source content
            podcast_mode: The podcast generation mode
        
        Returns:
            Calculated target word count for the podcast
        """
        content_words = len(content.split())
        
        if podcast_mode == "summary":
            # Quick overview with substance - ~17 minutes (~2500 words at 150 wpm speaking rate)
            return 2500
        elif podcast_mode == "analysis":
            # Deep discussion with thorough insights and examples
            return 4000
        else:  # "full" mode
            # Complete coverage - covers EVERYTHING, scales with content
            if content_words < 500:
                return max(1000, int(content_words * 2.0))
            elif content_words < 2000:
                return max(2000, int(content_words * 1.5))
            elif content_words < 5000:
                return max(3000, int(content_words * 1.2))
            else:
                # For very long content, cover everything comprehensively
                return min(10000, int(content_words * 1.0))
    
    def _get_mode_instructions(self, podcast_mode: str) -> str:
        """
        Get specific instructions for each podcast mode.
        
        Args:
            podcast_mode: The podcast generation mode
        
        Returns:
            Mode-specific instructions for the LLM
        """
        if podcast_mode == "summary":
            return """
SUMMARY MODE - Comprehensive Overview (~17 minutes):
- Cover the key points thoroughly, not just surface-level
- Explain the "why" behind each point, not just the "what"
- Each speaker should say 3-5 sentences per turn minimum
- Include relevant context and background for each major point
- The Host should ask follow-up questions to dig deeper into important concepts
- The Expert should provide concrete examples and real-world applications
- Total output: ~2500 words minimum
- Think: "What would a thorough 17-minute podcast that actually teaches the listener something cover?"
- DO NOT rush through topics - take time to explain concepts properly
- When the Expert explains something, the Host should often ask "Can you give an example?" or "Why does that matter?"
"""
        elif podcast_mode == "analysis":
            return """
ANALYSIS MODE - Deep Dive Discussion:
- Go DEEP into every topic - this is an in-depth analysis, not a surface overview
- For each major point, explore: the underlying principles, real-world implications, common misconceptions, and practical applications
- Each speaker should say 5-8 sentences per turn minimum
- Include detailed examples, case studies, and analogies
- The Host should challenge the Expert's assertions and ask probing questions
- The Expert should provide nuanced analysis with multiple perspectives
- Total output: ~4000 words minimum
- Think: "What would a 25+ minute deep-dive podcast between two experts cover?"
- DO NOT skip nuances or complexities - embrace them
- When discussing implications, go beyond the obvious - explore second and third-order effects
"""
        else:  # "full" mode
            return """
FULL MODE - Exhaustive Coverage:
- Cover EVERY SINGLE POINT in the content - nothing should be left out
- This is the most comprehensive mode - treat it like writing a book about the content
- Each speaker should say 5-10 sentences per turn minimum
- For each point: explain it, give context, provide examples, discuss implications, address counterarguments
- The Host should ensure thorough coverage by asking detailed follow-up questions
- The Expert should provide comprehensive analysis with extensive examples
- Total output: Scale based on content length (see word count target above)
- Think: "If someone could NOT read the original content, would this podcast teach them everything?"
- DO NOT summarize or condense - expand and elaborate on everything
- Include tangential but relevant information that enriches understanding
- When the content mentions something briefly, take time to fully explore it
"""
    
    def generate_script(self, content: str, conversation_config: dict) -> str:
        """
        Generate a podcast script from content using LM Studio.
        
        Args:
            content: The source content to discuss
            conversation_config: Conversation configuration dict
        
        Returns:
            Generated script text
        """
        podcast_mode = conversation_config.get('podcast_mode', 'summary')
        podcast_name = conversation_config.get('podcast_name', 'AI Podcast')
        creativity = conversation_config.get('creativity', 0.7)
        user_instructions = conversation_config.get('user_instructions', '')
        
        target_word_count = self._calculate_word_count(content, podcast_mode)
        mode_instructions = self._get_mode_instructions(podcast_mode)
        
        system_prompt = f"""You are a script writer for "{podcast_name}", a popular educational podcast. 
Your job is to take source content and transform it into an engaging, natural-sounding podcast dialogue.

CREATIVITY LEVEL: {creativity} (0.0=factual/strict, 1.0=creative/liberal interpretation)

TARGET LENGTH: Approximately {target_word_count} words total (both speakers combined).
This is CRITICAL - aim for this target length. Do not fall significantly short.

{mode_instructions}

OUTPUT FORMAT:
- Write as a dialogue between two speakers: "Host:" and "Expert:"
- Each line should start with the speaker name followed by a colon
- Alternate between speakers naturally (like a real conversation)
- Include natural speech patterns: filler words (um, uh, you know), contractions, pauses (...)
- Make it sound like a REAL conversation, not a scripted reading
- The Host is curious, asks good questions, and helps guide the discussion
- The Expert provides deep insights, examples, and explanations

NATURAL SPEECH GUIDELINES:
- Use contractions (don't, can't, it's, we're, etc.)
- Include occasional filler words naturally (um, uh, well, you know, I mean)
- Use pauses (...) for thinking or emphasis
- React naturally ("Oh wow", "That's interesting", "Wait, really?")
- Interrupt occasionally or build on each other's points
- Use casual transitions ("So basically", "In other words", "Here's the thing")

STYLE GUIDELINES:
- Conversational and engaging, not dry or academic
- Use analogies and real-world examples to explain complex concepts
- Include brief personal anecdotes or references where appropriate
- Balance depth with accessibility
- Keep the energy up - this should be fun to listen to!
"""
        
        if user_instructions:
            system_prompt += f"\n\nADDITIONAL INSTRUCTIONS FROM USER:\n{user_instructions}"
        
        user_message = f"""Please create a podcast script based on the following content.

Remember:
- Target length: ~{target_word_count} words
- Format: "Host: ..." and "Expert: ..." alternating
- Make it sound like a REAL, natural conversation
- Cover the content thoroughly according to {podcast_mode} mode

CONTENT TO CONVERT TO PODCAST:

{content}
"""
        
        logger.info(f"🧠 Generating podcast script ({podcast_mode} mode, ~{target_word_count} words)...")
        logger.info(f"   Content length: {len(content)} chars ({len(content.split())} words)")
        
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_message}
                ],
                temperature=creativity,
                max_tokens=self.config.llm.max_tokens,
            )
            
            script = response.choices[0].message.content
            logger.info(f"✅ Script generated: {len(script)} chars ({len(script.split())} words)")
            return script
            
        except Exception as e:
            logger.error(f"Failed to generate script: {e}")
            raise
    
    def generate_coverage_script(self, content: str) -> str:
        """
        Generate a script coverage report from a screenplay/CCSL using LM Studio.
        
        Uses the coverage system prompt from config. The output contains two parts:
        1. A structured markdown coverage report (for text reference)
        2. An audio narration script with "Analyst:" prefixes (for TTS)
        
        Args:
            content: The screenplay, script, or CCSL content
        
        Returns:
            Generated coverage text (both structured report and narration script)
        """
        coverage_config = self.config.coverage
        system_prompt = coverage_config.system_prompt
        
        if not system_prompt:
            raise ValueError(
                "Coverage system prompt is not configured. "
                "Please add a 'coverage.system_prompt' section to config.yaml."
            )
        
        # Inject analyst name into the prompt
        analyst_name = coverage_config.analyst_name
        system_prompt = system_prompt.replace("[Analyst Name]", analyst_name)
        
        from datetime import datetime
        today = datetime.now().strftime("%m/%d/%y")
        system_prompt = system_prompt.replace("[Today's Date in MM/DD/YY format]", today)
        
        user_message = f"""Please generate a complete script coverage report for the following screenplay/CCSL.

Output BOTH parts as specified:
1. Part 1: The structured markdown coverage report (between ---COVERAGE REPORT--- and ---END REPORT--- markers)
2. Part 2: The full audio narration script (with "Analyst:" prefixed sections)

Do NOT abbreviate or skip any sections. The narration script must cover everything in the structured report.

SCREENPLAY/CCSL CONTENT:

{content}
"""
        
        logger.info("🧠 Generating script coverage...")
        logger.info(f"   Content length: {len(content)} chars ({len(content.split())} words)")
        logger.info(f"   Analyst: {analyst_name}")
        
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_message}
                ],
                temperature=self.config.llm.temperature,
                max_tokens=self.config.llm.max_tokens,
            )
            
            script = response.choices[0].message.content
            logger.info(f"✅ Coverage generated: {len(script)} chars ({len(script.split())} words)")
            return script
            
        except Exception as e:
            logger.error(f"Failed to generate coverage: {e}")
            raise


class KokoroTTSClient:
    """Client for Kokoro TTS via OpenAI-compatible API."""
    
    def __init__(self, config: Config):
        self.config = config
        self.client = None
        self.server = None
    
    def start(self):
        """Start the Kokoro TTS server and initialize client."""
        kokoro_config = self.config.tts.kokoro
        
        # Check if server is already running
        if is_port_in_use(kokoro_config.server.host, kokoro_config.server.port):
            logger.info(f"🔗 Kokoro TTS server already running on port {kokoro_config.server.port}")
        elif kokoro_config.server.auto_start:
            # Start embedded server
            self.server = EmbeddedTTSServer(
                host=kokoro_config.server.host,
                port=kokoro_config.server.port
            )
            self.server.start()
            logger.info(f"🚀 Started embedded Kokoro TTS server on port {kokoro_config.server.port}")
            time.sleep(3)  # Wait for server to initialize
        else:
            raise RuntimeError(
                f"Kokoro TTS server not running on port {kokoro_config.server.port}. "
                f"Start it with: python src/tts_server.py"
            )
        
        # Create OpenAI client pointing to TTS server
        self.client = OpenAI(
            base_url=f"http://{kokoro_config.server.host}:{kokoro_config.server.port}/v1",
            api_key="not-needed"
        )
    
    def synthesize(self, text: str, voice: str) -> bytes:
        """
        Synthesize text to audio using Kokoro TTS.
        
        Args:
            text: Text to synthesize
            voice: Voice name (e.g., "af_bella")
        
        Returns:
            Audio bytes (WAV format)
        """
        if not self.client:
            raise RuntimeError("Kokoro TTS client not initialized. Call start() first.")
        
        try:
            response = self.client.audio.speech.create(
                model="kokoro",
                voice=voice,
                input=text
            )
            return response.content
        except Exception as e:
            logger.error(f"Kokoro TTS synthesis failed: {e}")
            raise
    
    def synthesize_batch(self, segments: List[PodcastSegment], max_workers: int = 4) -> List[bytes]:
        """
        Synthesize multiple segments in parallel.
        
        Args:
            segments: List of podcast segments
            max_workers: Number of parallel workers
        
        Returns:
            List of audio bytes in order
        """
        logger.info(f"🎙️ Synthesizing {len(segments)} audio segments with Kokoro TTS...")
        start_time = time.time()
        
        results = [None] * len(segments)
        
        def synthesize_one(index_seg):
            idx, seg = index_seg
            try:
                audio = self.synthesize(seg.text, seg.voice)
                return idx, audio
            except Exception as e:
                logger.error(f"Failed to synthesize segment {idx}: {e}")
                raise
        
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            futures = {executor.submit(synthesize_one, (i, seg)): i for i, seg in enumerate(segments)}
            
            completed = 0
            for future in as_completed(futures):
                idx, audio = future.result()
                results[idx] = audio
                completed += 1
                if completed % 5 == 0:
                    logger.info(f"   Progress: {completed}/{len(segments)} segments")
        
        elapsed = time.time() - start_time
        logger.info(f"✅ Kokoro audio synthesis complete in {elapsed:.1f}s")
        
        return results


class VoiceCloneTTSClient:
    """Client for Qwen3-TTS voice cloning."""
    
    def __init__(self, config: Config):
        self.config = config
        self.model = None
        self.voice_prompts = {}
    
    def start(self):
        """Load the voice cloning model and create voice prompts."""
        vc_config = self.config.tts.voice_clone
        
        logger.info("🧠 Loading Qwen3-TTS voice cloning model...")
        
        try:
            from qwen_tts import Qwen3TTSModel
            
            # Determine device
            device = vc_config.device
            if device == "auto":
                import torch
                if torch.cuda.is_available():
                    device = "cuda:0"
                elif hasattr(__import__('torch'), 'backends') and __import__('torch').backends.mps.is_available():
                    device = "mps"
                else:
                    device = "cpu"
            
            logger.info(f"   Device: {device}")
            
            self.model = Qwen3TTSModel.from_pretrained(vc_config.model)
            logger.info("✅ Model loaded successfully")
            
            # Create voice prompts for each speaker
            self._create_voice_prompts()
            
        except ImportError:
            raise ImportError(
                "qwen_tts package not installed. "
                "Install with: pip install qwen-tts"
            )
        except Exception as e:
            logger.error(f"Failed to load voice cloning model: {e}")
            raise
    
    def _create_voice_prompts(self):
        """Create voice clone prompts from reference audio files."""
        vc_config = self.config.tts.voice_clone
        
        for speaker_id in ["speaker_1", "speaker_2"]:
            speaker_config = vc_config.voices.__dict__.get(speaker_id)
            if not speaker_config:
                continue
            
            ref_audio = self._resolve_path(speaker_config.ref_audio)
            ref_text = speaker_config.ref_text
            
            logger.info(f"🎤 Creating voice prompt for {speaker_id} ({speaker_config.profile})...")
            
            voice_prompt = self.model.create_voice_clone_prompt(
                ref_audio=ref_audio,
                ref_text=ref_text
            )
            self.voice_prompts[speaker_id] = voice_prompt
            logger.info(f"   ✅ Voice prompt created for {speaker_config.profile}")
    
    def _resolve_path(self, path: str) -> str:
        """Resolve path relative to config file location."""
        config_path = Path(self.config.output.directory).parent
        resolved = config_path / path
        if resolved.exists():
            return str(resolved)
        # Try absolute path
        if Path(path).exists():
            return path
        raise FileNotFoundError(f"Reference audio not found: {path}")
    
    def _preprocess_text_for_conversation(self, text: str) -> str:
        """
        Preprocess text to make it sound more conversational for voice cloning.
        
        This adds natural speech patterns that help the TTS model produce
        more natural-sounding output instead of flat "reading" style.
        
        Args:
            text: Original text to preprocess
        
        Returns:
            Preprocessed text with conversational markers
        """
        import re
        
        # Add brief pauses after punctuation for natural rhythm
        text = re.sub(r'\.(\s|$)', '. ', text)
        text = re.sub(r'\?(\s|$)', '? ', text)
        text = re.sub(r'\!(\s|$)', '! ', text)
        
        # Add comma pauses for natural breathing
        text = re.sub(r',(\s)', ', ', text)
        
        # Handle ellipsis for thinking pauses
        text = re.sub(r'\.\.\.', '... ', text)
        
        # Add slight pause before certain words for emphasis
        emphasis_words = ['however', 'but', 'actually', 'in fact', 'you know', 'so', 'well']
        for word in emphasis_words:
            text = re.sub(rf'\b{word}\b', f'... {word}', text, count=1)
        
        # Clean up any double spaces
        text = re.sub(r'\s{2,}', ' ', text)
        
        return text.strip()
    
    def synthesize(self, text: str, speaker: str) -> bytes:
        """
        Synthesize text to audio using voice cloning.
        
        Args:
            text: Text to synthesize
            speaker: Speaker identifier ("speaker_1" or "speaker_2")
        
        Returns:
            Audio bytes (WAV format)
        """
        if not self.model:
            raise RuntimeError("Voice clone TTS client not initialized. Call start() first.")
        
        voice_prompt = self.voice_prompts.get(speaker)
        if not voice_prompt:
            raise RuntimeError(f"No voice prompt for speaker: {speaker}")
        
        vc_config = self.config.tts.voice_clone
        speaker_config = vc_config.voices.__dict__.get(speaker)
        language = speaker_config.language if speaker_config else "English"
        
        # Preprocess text for more conversational output
        processed_text = self._preprocess_text_for_conversation(text)
        
        try:
            import soundfile as sf
            import io
            
            wavs, sr = self.model.generate_voice_clone(
                text=processed_text,
                language=language,
                voice_clone_prompt=voice_prompt,
                max_new_tokens=vc_config.max_tokens,
            )
            
            if wavs:
                # Convert to WAV bytes
                buffer = io.BytesIO()
                sf.write(buffer, wavs[0], sr, format='WAV')
                buffer.seek(0)
                return buffer.read()
            else:
                raise RuntimeError("No audio generated")
                
        except Exception as e:
            logger.error(f"Voice clone TTS synthesis failed: {e}")
            raise
    
    def synthesize_batch(self, segments: List[PodcastSegment], max_workers: int = 1) -> List[bytes]:
        """
        Synthesize multiple segments sequentially (voice cloning doesn't support parallel well).
        
        Args:
            segments: List of podcast segments
            max_workers: Ignored for voice cloning (sequential processing)
        
        Returns:
            List of audio bytes in order
        """
        logger.info(f"🎙️ Synthesizing {len(segments)} audio segments with Voice Cloning...")
        start_time = time.time()
        
        results = []
        
        # Voice cloning is memory-intensive, process sequentially
        for i, seg in enumerate(segments):
            try:
                # Map speaker name to speaker_id
                # "Host" and "Analyst" both use speaker_1 (single narrator or host voice)
                speaker_lower = seg.speaker.lower()
                if speaker_lower in ("host", "analyst"):
                    speaker_id = "speaker_1"
                else:
                    speaker_id = "speaker_2"
                
                audio = self.synthesize(seg.text, speaker_id)
                results.append(audio)
                
                if (i + 1) % 5 == 0:
                    logger.info(f"   Progress: {i + 1}/{len(segments)} segments")
                    
            except Exception as e:
                logger.error(f"Failed to synthesize segment {i}: {e}")
                raise
        
        elapsed = time.time() - start_time
        logger.info(f"✅ Voice clone audio synthesis complete in {elapsed:.1f}s")
        
        return results


class TTSClient:
    """Factory for TTS clients based on provider configuration."""
    
    def __init__(self, config: Config):
        self.config = config
        self._client = None
    
    def _get_client(self):
        """Get the appropriate TTS client based on provider."""
        if self._client is None:
            provider = self.config.tts.provider
            if provider == "kokoro":
                self._client = KokoroTTSClient(self.config)
            elif provider == "voice_clone":
                self._client = VoiceCloneTTSClient(self.config)
            else:
                raise ValueError(f"Unknown TTS provider: {provider}")
        return self._client
    
    def start_server(self):
        """Start the TTS server/client (legacy method name)."""
        client = self._get_client()
        client.start()
    
    def synthesize(self, text: str, voice: str) -> bytes:
        """Synthesize text to audio."""
        client = self._get_client()
        return client.synthesize(text, voice)
    
    def synthesize_batch(self, segments: List[PodcastSegment], max_workers: int = 4) -> List[bytes]:
        """Synthesize multiple segments."""
        client = self._get_client()
        return client.synthesize_batch(segments, max_workers)


def parse_script(script: str, config: Config) -> List[PodcastSegment]:
    """
    Parse a script into segments.
    
    Args:
        script: Raw script text
        config: Configuration object (used to determine TTS provider and voice settings)
    
    Returns:
        List of PodcastSegment objects
    """
    segments = []
    lines = script.strip().split('\n')
    
    # Get voice info based on provider
    provider = config.tts.provider
    
    if provider == "voice_clone":
        # For voice cloning, voice field stores the speaker_id for lookup
        speaker_1_voice = "speaker_1"
        speaker_2_voice = "speaker_2"
    else:  # kokoro
        # For Kokoro, voice field stores the voice name
        speaker_1_voice = config.tts.kokoro.voices.speaker_1
        speaker_2_voice = config.tts.kokoro.voices.speaker_2
    
    # Track parsing stats for debugging
    total_lines = len(lines)
    matched_lines = 0
    skipped_lines = 0
    
    # More flexible regex patterns to handle various LLM output formats
    patterns = [
        r'^\*?\*?(Host|Expert|Speaker\s*1|Speaker\s*2)\*?\*?\s*:\s*(.+)$',
        r'^(HOST|EXPERT)\s*:\s*(.+)$',
        r'^(Speaker\s*[12])\s*:\s*(.+)$',
    ]
    
    for line in lines:
        line = line.strip()
        if not line:
            skipped_lines += 1
            continue
        
        # Skip markdown headers, titles, and non-dialogue lines
        if line.startswith('#') or line.startswith('---') or line.startswith('==='):
            skipped_lines += 1
            continue
        
        matched = False
        for pattern in patterns:
            match = re.match(pattern, line, re.IGNORECASE)
            if match:
                matched = True
                matched_lines += 1
                speaker = match.group(1).lower().replace('*', '').strip()
                text = match.group(2).strip()
                
                # Remove markdown formatting from text
                text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # Remove bold
                text = re.sub(r'\*(.+?)\*', r'\1', text)      # Remove italic
                
                # Map speaker to voice
                if speaker in ('host', 'speaker 1', 'speaker1', 'speaker1'):
                    voice = speaker_1_voice
                    speaker_name = "Host"
                else:
                    voice = speaker_2_voice
                    speaker_name = "Expert"
                
                if text:  # Only add non-empty segments
                    segments.append(PodcastSegment(
                        speaker=speaker_name,
                        text=text,
                        voice=voice
                    ))
                break
        
        if not matched:
            skipped_lines += 1
            if skipped_lines <= 5:
                logger.debug(f"   Skipped line: '{line[:80]}...'")
    
    # Log parsing stats
    logger.info(f"   Script parsing: {total_lines} lines, {matched_lines} matched, {skipped_lines} skipped")
    
    if not segments:
        logger.warning("⚠️ No dialogue segments parsed from script!")
        logger.warning("   Script preview (first 500 chars):")
        logger.warning(f"   {script[:500]}")
        logger.warning("   Expected format: 'Host: text' or 'Expert: text'")
    
    return segments


def parse_coverage_script(script: str, config: Config) -> CoverageResult:
    """
    Parse a coverage script into narration segments and structured report.
    
    The LLM output has two parts:
    1. Structured report between ---COVERAGE REPORT--- and ---END REPORT--- markers
    2. Narration script with "Analyst:" prefixed lines
    
    Args:
        script: Raw coverage text from LLM
        config: Configuration object
    
    Returns:
        CoverageResult with segments, structured report, and raw text
    """
    # Get voice for the analyst (always speaker_1)
    provider = config.tts.provider
    if provider == "voice_clone":
        analyst_voice = "speaker_1"
    else:
        analyst_voice = config.tts.kokoro.voices.speaker_1
    
    # Extract structured report
    structured_report = ""
    report_match = re.search(
        r'---\s*COVERAGE REPORT\s*---\s*\n(.*?)\n\s*---\s*END REPORT\s*---',
        script,
        re.DOTALL | re.IGNORECASE
    )
    if report_match:
        structured_report = report_match.group(1).strip()
        logger.info(f"   Extracted structured report: {len(structured_report)} chars")
    else:
        logger.warning("   Could not find structured report markers (---COVERAGE REPORT--- / ---END REPORT---)")
        structured_report = script
    
    # Extract narration section (everything after ---END REPORT--- or entire script)
    narration_text = script
    if report_match:
        end_pos = report_match.end()
        narration_text = script[end_pos:]
    
    # Parse "Analyst:" prefixed lines into segments
    segments = []
    lines = narration_text.strip().split('\n')
    current_text_parts = []
    
    for line in lines:
        stripped = line.strip()
        
        # Check if this line starts a new Analyst segment
        analyst_match = re.match(r'^\*?\*?Analyst\*?\*?\s*:\s*(.*)$', stripped, re.IGNORECASE)
        
        if analyst_match:
            # Save any accumulated text from previous segment
            if current_text_parts:
                full_text = ' '.join(current_text_parts).strip()
                # Remove markdown formatting
                full_text = re.sub(r'\*\*(.+?)\*\*', r'\1', full_text)
                full_text = re.sub(r'\*(.+?)\*', r'\1', full_text)
                if full_text:
                    segments.append(PodcastSegment(
                        speaker="Analyst",
                        text=full_text,
                        voice=analyst_voice
                    ))
            # Start new segment with text after "Analyst:"
            remainder = analyst_match.group(1).strip()
            current_text_parts = [remainder] if remainder else []
        elif stripped and not stripped.startswith('#'):
            # Continuation of current segment (non-empty, non-header lines)
            if current_text_parts is not None:
                current_text_parts.append(stripped)
    
    # Don't forget the last segment
    if current_text_parts:
        full_text = ' '.join(current_text_parts).strip()
        full_text = re.sub(r'\*\*(.+?)\*\*', r'\1', full_text)
        full_text = re.sub(r'\*(.+?)\*', r'\1', full_text)
        if full_text:
            segments.append(PodcastSegment(
                speaker="Analyst",
                text=full_text,
                voice=analyst_voice
            ))
    
    logger.info(f"   Coverage parsing: {len(segments)} narration segments extracted")
    
    if not segments:
        logger.warning("⚠️ No narration segments parsed from coverage!")
        logger.warning("   Attempting fallback: treating entire output as single narration...")
        # Fallback: try to find any paragraph-like blocks and use them
        paragraphs = re.split(r'\n\s*\n', narration_text.strip())
        for para in paragraphs:
            para = para.strip()
            if len(para) > 50:  # Skip very short fragments
                # Remove markdown formatting
                para = re.sub(r'\*\*(.+?)\*\*', r'\1', para)
                para = re.sub(r'\*(.+?)\*', r'\1', para)
                para = re.sub(r'^#+\s*', '', para, flags=re.MULTILINE)  # Remove headers
                if para:
                    segments.append(PodcastSegment(
                        speaker="Analyst",
                        text=para,
                        voice=analyst_voice
                    ))
        logger.info(f"   Fallback parsing: {len(segments)} segments from paragraph splitting")
    
    return CoverageResult(
        segments=segments,
        structured_report=structured_report,
        raw_text=script
    )


def load_content(input_path: str) -> str:
    """
    Load content from a file.
    
    Supports:
    - Native text formats: .md, .txt, .markdown (read directly)
    - Document formats: .pdf, .docx, .pptx, .doc, .ppt, .xlsx, .xls, .html (converted to markdown)
    
    Args:
        input_path: Path to input file
    
    Returns:
        Content as string (markdown format)
    """
    from .doc_converter import load_document_content
    
    return load_document_content(input_path)


def save_script(script: str, output_path: str):
    """Save script to file."""
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, 'w', encoding='utf-8') as f:
        f.write(script)


def save_coverage_docx(structured_report: str, output_path: str):
    """
    Save coverage report as a formatted Microsoft Word document.
    
    Parses the markdown-structured report and creates a properly formatted .docx file
    with tables, headings, bold text, and paragraphs.
    
    Args:
        structured_report: Markdown-formatted coverage report text
        output_path: Path for output .docx file
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        logger.warning("python-docx not installed. Skipping .docx generation. Install with: pip install python-docx")
        return
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    lines = structured_report.strip().split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Check for table (starts with |)
        if line.startswith('|'):
            # Collect all table rows
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            
            # Parse table
            if len(table_lines) >= 2:
                # Parse header row
                header_cells = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
                
                # Skip separator row (contains ---)
                data_start = 1
                if len(table_lines) > 1 and '---' in table_lines[1]:
                    data_start = 2
                
                # Create table
                num_cols = len(header_cells)
                num_rows = len(table_lines) - data_start
                table = doc.add_table(rows=num_rows + 1, cols=num_cols)
                table.style = 'Table Grid'
                
                # Fill header
                for j, cell_text in enumerate(header_cells):
                    cell = table.rows[0].cells[j]
                    cell.text = cell_text
                    # Make header bold
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                # Fill data rows
                for row_idx in range(data_start, len(table_lines)):
                    data_cells = [cell.strip() for cell in table_lines[row_idx].split('|')[1:-1]]
                    for col_idx, cell_text in enumerate(data_cells):
                        if col_idx < num_cols:
                            cell = table.rows[row_idx - data_start + 1].cells[col_idx]
                            # Remove markdown formatting
                            cell_text = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)
                            cell.text = cell_text
                
                doc.add_paragraph()  # Spacing after table
            continue
        
        # Check for heading (LOGLINE, SUMMARY, COMMENTS)
        if line.startswith('**') and line.endswith(':**'):
            heading_text = line.strip('*:')
            heading = doc.add_heading(heading_text, level=2)
            i += 1
            continue
        
        # Check for recommendation line (bold RECOMMEND/CONSIDER/PASS)
        if line.startswith('**') and ('RECOMMEND' in line or 'CONSIDER' in line or 'PASS' in line):
            rec_text = line.strip('*')
            para = doc.add_paragraph()
            run = para.add_run(rec_text)
            run.bold = True
            run.font.size = Pt(12)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        
        # Regular paragraph
        # Collect consecutive non-empty lines as one paragraph
        para_lines = []
        while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith('|') and not lines[i].strip().startswith('**'):
            para_lines.append(lines[i].strip())
            i += 1
        
        if para_lines:
            para_text = ' '.join(para_lines)
            # Handle bold text within paragraph
            para = doc.add_paragraph()
            
            # Split by bold markers and add runs
            parts = re.split(r'(\*\*.*?\*\*)', para_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    # Bold text
                    run = para.add_run(part[2:-2])
                    run.bold = True
                else:
                    # Regular text
                    para.add_run(part)
    
    # Save document
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    logger.info(f"📄 Coverage report saved as Word document: {output_path}")


def generate_silence(duration_seconds: float, sample_rate: int = 24000) -> np.ndarray:
    """
    Generate silence audio samples.
    
    Args:
        duration_seconds: Duration of silence in seconds
        sample_rate: Sample rate for the audio
    
    Returns:
        numpy array of silence samples
    """
    import numpy as np
    num_samples = int(duration_seconds * sample_rate)
    return np.zeros(num_samples, dtype=np.int16)


def concatenate_audio(audio_chunks: List[bytes], output_path: str,
                      segments: List[PodcastSegment] = None,
                      pause_between: float = 0.5,
                      pause_on_change: float = 1.0):
    """
    Concatenate audio chunks with natural pauses between dialogue segments.
    
    Args:
        audio_chunks: List of WAV audio bytes
        output_path: Output file path
        segments: List of podcast segments (to detect speaker changes)
        pause_between: Seconds of silence between each segment
        pause_on_change: Extra pause when speaker changes (turn-taking)
    """
    import scipy.io.wavfile as wavfile
    import io
    
    # Parse all WAV files and concatenate with pauses
    samples_list = []
    sample_rate = 24000  # Default
    
    prev_speaker = None
    
    logger.info(f"🔍 Debug: Received {len(audio_chunks)} audio chunks")
    for i, chunk in enumerate(audio_chunks):
        if chunk is None:
            logger.warning(f"   Chunk {i}: is None")
        elif len(chunk) == 0:
            logger.warning(f"   Chunk {i}: is empty bytes (0 bytes)")
        else:
            logger.debug(f"   Chunk {i}: {len(chunk)} bytes")
        
    for i, chunk in enumerate(audio_chunks):
        if chunk:
            buffer = io.BytesIO(chunk)
            sr, samples = wavfile.read(buffer)
            
            if not samples_list:
                sample_rate = sr
                samples_list.append(samples)
            else:
                # Add pause before this segment
                if segments and i < len(segments):
                    current_speaker = segments[i].speaker
                    
                    # Determine pause duration based on speaker change
                    if prev_speaker and current_speaker != prev_speaker:
                        # Speaker changed - add longer pause (turn-taking)
                        pause_duration = pause_between + pause_on_change
                    else:
                        # Same speaker continuing - shorter pause
                        pause_duration = pause_between
                    
                    silence = generate_silence(pause_duration, sample_rate)
                    samples_list.append(silence)
                    prev_speaker = current_speaker
                else:
                    # No segment info - use default pause
                    silence = generate_silence(pause_between, sample_rate)
                    samples_list.append(silence)
                
                samples_list.append(samples)
    
    if not samples_list:
        # Provide more detailed error message
        none_count = sum(1 for c in audio_chunks if c is None)
        empty_count = sum(1 for c in audio_chunks if c is not None and len(c) == 0)
        valid_count = sum(1 for c in audio_chunks if c is not None and len(c) > 0)
        error_msg = (
            f"No audio data to save. "
            f"Total chunks: {len(audio_chunks)}, "
            f"Valid: {valid_count}, "
            f"Empty: {empty_count}, "
            f"None: {none_count}"
        )
        logger.error(f"❌ {error_msg}")
        raise ValueError(error_msg)
    
    # Concatenate all samples
    combined = np.concatenate(samples_list)
    
    # Save to file
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    
    wavfile.write(path, sample_rate, combined)
    logger.info(f"💾 Audio saved to: {output_path}")


class PodcastGenerator:
    """
    Main podcast generator that orchestrates the entire pipeline.
    """
    
    def __init__(self, config: Optional[Config] = None):
        self.config = config or get_config()
        self.llm_client = LLMClient(self.config)
        self.tts_client = TTSClient(self.config)
    
    def generate(self, input_path: str, output_path: str, 
                 script_only: bool = False) -> str:
        """
        Generate a podcast from input content.
        
        Args:
            input_path: Path to input content file
            output_path: Path for output audio file
            script_only: If True, only generate script without audio
        
        Returns:
            Path to generated file (audio or script)
        """
        # Check if this is coverage mode
        if self.config.conversation.podcast_mode == "coverage":
            return self.generate_coverage(input_path, output_path, script_only)
        
        # Load content
        logger.info(f"📂 Loading content from: {input_path}")
        content = load_content(input_path)
        logger.info(f"   Content length: {len(content)} characters")
        
        # Check LLM connection
        if not self.llm_client.check_connection():
            raise RuntimeError(
                f"Cannot connect to LM Studio at {self.config.llm.base_url}. "
                "Please ensure LM Studio is running with the model loaded."
            )
        
        # Generate script
        script_text = self.llm_client.generate_script(
            content, 
            vars(self.config.conversation)
        )
        
        # Parse script into segments
        segments = parse_script(script_text, self.config)
        logger.info(f"📝 Parsed {len(segments)} dialogue segments")
        
        # Check if we have any segments to synthesize
        if not segments:
            raise RuntimeError(
                "No dialogue segments found in the generated script. "
                "The LLM may have returned a non-dialogue format. "
                "Check the script output and ensure it uses 'Host:' and 'Expert:' format."
            )
        
        # Save script if requested
        script_path = Path(output_path).with_suffix('.txt')
        if self.config.output.keep_temp_script or script_only:
            save_script(script_text, str(script_path))
            logger.info(f"📄 Script saved to: {script_path}")
        
        if script_only:
            return str(script_path)
        
        # Start TTS server
        self.tts_client.start_server()
        
        # Synthesize audio
        audio_chunks = self.tts_client.synthesize_batch(segments)
        
        # Concatenate and save with natural pauses
        concatenate_audio(
            audio_chunks,
            output_path,
            segments=segments,
            pause_between=self.config.tts.pause_between_segments,
            pause_on_change=self.config.tts.pause_on_speaker_change
        )
        
        return output_path
    
    def generate_from_text(self, text: str, output_path: str,
                          script_only: bool = False) -> str:
        """
        Generate a podcast from raw text.
        
        Args:
            text: Raw text content
            output_path: Path for output audio file
            script_only: If True, only generate script without audio
        
        Returns:
            Path to generated file
        """
        # Check if this is coverage mode
        if self.config.conversation.podcast_mode == "coverage":
            return self.generate_coverage_from_text(text, output_path, script_only)
        
        # Check LLM connection
        if not self.llm_client.check_connection():
            raise RuntimeError(
                f"Cannot connect to LM Studio at {self.config.llm.base_url}. "
                "Please ensure LM Studio is running with the model loaded."
            )
        
        # Generate script
        script_text = self.llm_client.generate_script(
            text,
            vars(self.config.conversation)
        )
        
        # Parse script into segments
        segments = parse_script(script_text, self.config)
        logger.info(f"📝 Parsed {len(segments)} dialogue segments")
        
        # Check if we have any segments to synthesize
        if not segments:
            raise RuntimeError(
                "No dialogue segments found in the generated script. "
                "The LLM may have returned a non-dialogue format. "
                "Check the script output and ensure it uses 'Host:' and 'Expert:' format."
            )
        
        # Save script if requested
        script_path = Path(output_path).with_suffix('.txt')
        if self.config.output.keep_temp_script or script_only:
            save_script(script_text, str(script_path))
            logger.info(f"📄 Script saved to: {script_path}")
        
        if script_only:
            return str(script_path)
        
        # Start TTS server
        self.tts_client.start_server()
        
        # Synthesize audio
        audio_chunks = self.tts_client.synthesize_batch(segments)
        
        # Concatenate and save with natural pauses
        concatenate_audio(
            audio_chunks,
            output_path,
            segments=segments,
            pause_between=self.config.tts.pause_between_segments,
            pause_on_change=self.config.tts.pause_on_speaker_change
        )
        
        return output_path
    
    def generate_coverage(self, input_path: str, output_path: str,
                          script_only: bool = False) -> str:
        """
        Generate a script coverage audio from a screenplay/CCSL.
        
        Single-speaker narration of a structured coverage report.
        
        Args:
            input_path: Path to screenplay/CCSL file
            output_path: Path for output audio file
            script_only: If True, only generate the report without audio
        
        Returns:
            Path to generated file (audio or report)
        """
        # Load content
        logger.info(f"📂 Loading screenplay/CCSL from: {input_path}")
        content = load_content(input_path)
        logger.info(f"   Content length: {len(content)} characters")
        
        # Check LLM connection
        if not self.llm_client.check_connection():
            raise RuntimeError(
                f"Cannot connect to LM Studio at {self.config.llm.base_url}. "
                "Please ensure LM Studio is running with the model loaded."
            )
        
        # Generate coverage
        coverage_text = self.llm_client.generate_coverage_script(content)
        
        # Parse coverage into segments and structured report
        coverage_result = parse_coverage_script(coverage_text, self.config)
        logger.info(f"📝 Parsed {len(coverage_result.segments)} narration segments")
        
        # Always save the structured report
        report_path = Path(output_path).with_suffix('.txt')
        save_script(coverage_result.structured_report, str(report_path))
        logger.info(f"📄 Coverage report saved to: {report_path}")
        
        # Also save as Word document
        docx_path = Path(output_path).with_suffix('.docx')
        save_coverage_docx(coverage_result.structured_report, str(docx_path))
        
        # Check if we have any segments to synthesize
        if not coverage_result.segments:
            raise RuntimeError(
                "No narration segments found in the generated coverage. "
                "The LLM may not have produced 'Analyst:' prefixed sections. "
                "Check the coverage output format."
            )
        
        if script_only:
            return str(report_path)
        
        # Start TTS server
        self.tts_client.start_server()
        
        # Synthesize audio (single speaker)
        audio_chunks = self.tts_client.synthesize_batch(coverage_result.segments)
        
        # Concatenate with coverage-specific pauses (no speaker-change pauses needed)
        concatenate_audio(
            audio_chunks,
            output_path,
            segments=coverage_result.segments,
            pause_between=self.config.coverage.pause_between_segments,
            pause_on_change=0.0  # Single speaker, no turn-taking pauses
        )
        
        return output_path
    
    def generate_coverage_from_text(self, text: str, output_path: str,
                                    script_only: bool = False) -> str:
        """
        Generate a script coverage audio from raw text.
        
        Args:
            text: Raw screenplay/CCSL text
            output_path: Path for output audio file
            script_only: If True, only generate the report without audio
        
        Returns:
            Path to generated file
        """
        # Check LLM connection
        if not self.llm_client.check_connection():
            raise RuntimeError(
                f"Cannot connect to LM Studio at {self.config.llm.base_url}. "
                "Please ensure LM Studio is running with the model loaded."
            )
        
        # Generate coverage
        coverage_text = self.llm_client.generate_coverage_script(text)
        
        # Parse coverage into segments and structured report
        coverage_result = parse_coverage_script(coverage_text, self.config)
        logger.info(f"📝 Parsed {len(coverage_result.segments)} narration segments")
        
        # Always save the structured report
        report_path = Path(output_path).with_suffix('.txt')
        save_script(coverage_result.structured_report, str(report_path))
        logger.info(f"📄 Coverage report saved to: {report_path}")
        
        # Also save as Word document
        docx_path = Path(output_path).with_suffix('.docx')
        save_coverage_docx(coverage_result.structured_report, str(docx_path))
        
        # Check if we have any segments to synthesize
        if not coverage_result.segments:
            raise RuntimeError(
                "No narration segments found in the generated coverage. "
                "The LLM may not have produced 'Analyst:' prefixed sections. "
                "Check the coverage output format."
            )
        
        if script_only:
            return str(report_path)
        
        # Start TTS server
        self.tts_client.start_server()
        
        # Synthesize audio (single speaker)
        audio_chunks = self.tts_client.synthesize_batch(coverage_result.segments)
        
        # Concatenate with coverage-specific pauses
        concatenate_audio(
            audio_chunks,
            output_path,
            segments=coverage_result.segments,
            pause_between=self.config.coverage.pause_between_segments,
            pause_on_change=0.0
        )
        
        return output_path
